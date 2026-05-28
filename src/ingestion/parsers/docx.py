"""DOCX document parser."""
import logging
from pathlib import Path
from typing import Optional, List, Dict, Any
from datetime import datetime

from .base import DocumentParser, ParsedDocument, DocumentMetadata


class DOCXParser(DocumentParser):
    """
    Parse DOCX (Word) documents and extract text and metadata.

    Uses python-docx library to extract text while preserving formatting info.
    Handles tables, lists, and document structure.
    """

    def __init__(self):
        """Initialize the DOCX parser."""
        super().__init__()
        self.logger = logging.getLogger(__name__)

    def parse(self, file_path: Path) -> ParsedDocument:
        """
        Parse a DOCX file and extract content and metadata.

        Args:
            file_path: Path to the DOCX file

        Returns:
            ParsedDocument with extracted content and metadata

        Raises:
            IOError: If file cannot be read
            ValueError: If DOCX is invalid
        """
        if isinstance(file_path, str):
            file_path = Path(file_path)

        if not file_path.exists():
            raise IOError(f"DOCX file not found: {file_path}")

        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx not available. "
                "Install it: pip install python-docx"
            )

        try:
            doc = Document(file_path)

            content = self._extract_text(doc)
            metadata = self._extract_metadata_from_docx(doc, file_path)
            sections = self._extract_sections(doc)

            self.logger.info(
                f"Successfully parsed DOCX: {file_path.name}",
                extra={'paragraphs': len(doc.paragraphs)}
            )

            return ParsedDocument(
                content=content,
                metadata=metadata,
                sections=sections,
            )

        except Exception as e:
            self.logger.error(
                f"Failed to parse DOCX {file_path}: {str(e)}"
            )
            raise ValueError(f"Invalid or corrupted DOCX file: {file_path}") from e

    def extract_metadata(self, file_path: Path) -> DocumentMetadata:
        """
        Extract metadata from a DOCX file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            DocumentMetadata with extracted metadata
        """
        if isinstance(file_path, str):
            file_path = Path(file_path)

        try:
            from docx import Document
        except ImportError:
            return DocumentMetadata(file_path=str(file_path))

        try:
            doc = Document(file_path)
            return self._extract_metadata_from_docx(doc, file_path)
        except Exception as e:
            self.logger.warning(
                f"Could not extract metadata from {file_path}: {e}"
            )
            return DocumentMetadata(file_path=str(file_path))

    def _extract_text(self, doc) -> str:
        """
        Extract all text from DOCX while preserving structure.

        Includes paragraphs, tables, and text boxes.

        Args:
            doc: Document object from python-docx

        Returns:
            Extracted text
        """
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            table_text = self._extract_table_text(table)
            if table_text:
                text_parts.append(table_text)

        # Extract text from sections if available
        for section in getattr(doc, 'sections', []):
            header_text = self._extract_section_header(section)
            footer_text = self._extract_section_footer(section)

            if header_text:
                text_parts.insert(0, header_text)
            if footer_text:
                text_parts.append(footer_text)

        return "\n\n".join(text_parts)

    def _extract_table_text(self, table) -> str:
        """
        Extract text from a table.

        Args:
            table: Table object from python-docx

        Returns:
            Text representation of table
        """
        rows = []

        for row in table.rows:
            cells = []
            for cell in row.cells:
                # Extract text from cell paragraphs
                cell_text = " ".join(
                    p.text for p in cell.paragraphs
                ).strip()
                cells.append(cell_text)

            if cells:
                rows.append(" | ".join(cells))

        return "\n".join(rows) if rows else ""

    def _extract_section_header(self, section) -> Optional[str]:
        """
        Extract header text from section.

        Args:
            section: Section object

        Returns:
            Header text or None
        """
        try:
            header = section.header
            if header:
                return "\n".join(p.text for p in header.paragraphs)
        except Exception:
            pass

        return None

    def _extract_section_footer(self, section) -> Optional[str]:
        """
        Extract footer text from section.

        Args:
            section: Section object

        Returns:
            Footer text or None
        """
        try:
            footer = section.footer
            if footer:
                return "\n".join(p.text for p in footer.paragraphs)
        except Exception:
            pass

        return None

    def _extract_metadata_from_docx(
        self, doc, file_path: Path
    ) -> DocumentMetadata:
        """
        Extract metadata from DOCX document properties.

        Args:
            doc: Document object from python-docx
            file_path: Path to DOCX file

        Returns:
            DocumentMetadata with extracted metadata
        """
        # Get core properties
        props = doc.core_properties

        title = props.title if props.title else None
        author = props.author if props.author else None
        subject = props.subject if props.subject else None

        # Extract dates
        created_date = props.created
        modified_date = props.modified

        # File info
        file_size = file_path.stat().st_size if file_path.exists() else None

        # Count content
        para_count = len(doc.paragraphs)
        table_count = len(doc.tables)

        # Estimate word count
        full_text = "\n".join(p.text for p in doc.paragraphs)
        word_count = len(full_text.split()) if full_text else None

        extra_meta: Dict[str, Any] = {}

        if subject:
            extra_meta['subject'] = subject

        if props.comments:
            extra_meta['comments'] = props.comments

        if props.category:
            extra_meta['category'] = props.category

        if props.keywords:
            extra_meta['keywords'] = props.keywords

        extra_meta['paragraph_count'] = para_count
        extra_meta['table_count'] = table_count

        return DocumentMetadata(
            title=title,
            author=author,
            created_date=created_date,
            modified_date=modified_date,
            word_count=word_count,
            file_size=file_size,
            file_path=str(file_path),
            extra=extra_meta,
        )

    def _extract_sections(self, doc) -> List[str]:
        """
        Extract document sections/headings.

        Args:
            doc: Document object from python-docx

        Returns:
            List of section headings
        """
        sections = []

        for para in doc.paragraphs:
            # Check if paragraph is a heading
            style = para.style
            if style and 'Heading' in style.name:
                if para.text.strip():
                    sections.append(para.text.strip())

        return sections[:20]  # Limit to first 20 sections
